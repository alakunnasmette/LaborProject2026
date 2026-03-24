# report_generator.py

from docx import Document
from docx.shared import Pt
from text_library import (
    static_texts,
    loopbaanfase_texts,
    personality_texts,
    loopbaanankers_texts,
    carriereclusters_texts,
    cultuuranalyse_texts
)

# -------------------- CONFIGURATION --------------------

JCM_MAPPING = {
    "Taakvaardigheid": ("SKILL_VARIETY_TITLE", "SKILL_VARIETY_TEXT"),
    "Taakidentiteit": ("TASK_IDENTITY_TITLE", "TASK_IDENTITY_TEXT"),
    "Taakbetekenis": ("TASK_SIGNIFICANCE_TITLE", "TASK_SIGNIFICANCE_TEXT"),
    "Autonomie": ("AUTONOMY_TITLE", "AUTONOMY_TEXT"),
    "Feedback": ("FEEDBACK_TITLE", "FEEDBACK_TEXT"),
}


LOOPBAANFASE_RANGES = [
    (17, 21, "LOOPBAANFASE_17_21"),
    (22, 28, "LOOPBAANFASE_22_28"),
    (29, 33, "LOOPBAANFASE_29_33"),
    (34, 40, "LOOPBAANFASE_34_40"),
    (41, 45, "LOOPBAANFASE_41_45"),
    (46, 64, "LOOPBAANFASE_46_64"),
]


# -------------------- HELPERS --------------------

def add_bold_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_paragraph_conditional_indent(doc, line, left_indent_bullets=1.0, left_indent_normal=0.0, space_after=6):
    """
    Adds a paragraph with conditional indentation:
    - Bullet points ('•')
    """
    line = line.strip()
    if not line:
        return

    p = doc.add_paragraph(line)
    if line.startswith("•"):
        # Indent bullet points and short:long lines
        p.paragraph_format.left_indent = Pt(left_indent_bullets * 28.35)
    else:
        p.paragraph_format.left_indent = Pt(left_indent_normal * 28.35)
    p.paragraph_format.space_after = Pt(space_after)

def add_title_and_body(doc, text):
    lines = text.strip().split("\n")

    title = lines[0]
    body_lines = lines[1:]

    # Bold title
    add_bold_paragraph(doc, title)

    # Add each body paragraph separately
    for line in body_lines:
        line = line.strip()
        if line:
            doc.add_paragraph(line)


# -------------------- MAIN REPORT GENERATOR --------------------

def generate_report(user_data, filename):

    doc = Document()

    add_loopbaanfase_section(doc, user_data)
    add_bigfive_section(doc, user_data)
    add_loopbaanankers_section(doc, user_data)
    add_carriereclusters_section(doc, user_data)
    add_cultuur_section(doc, user_data)
    add_jcm_section(doc, user_data)

    doc.save(filename)


# -------------------- 1.0 LOOPBAANFASE --------------------

def add_loopbaanfase_section(doc, user_data):

    doc.add_heading(static_texts.LOOPBAANFASE_SECTION_TITLE, level=1)
    doc.add_paragraph(static_texts.LOOPBAANFASE_INTRO_TEXT)

    age = user_data.get("age", 0)

    text = loopbaanfase_texts.LOOPBAANFASE_60_OLDER

    for min_age, max_age, attr in LOOPBAANFASE_RANGES:
        if min_age <= age <= max_age:
            text = getattr(loopbaanfase_texts, attr)
            break

    lines = text.strip().split("\n")

    title = lines[0]
    body = "\n".join(lines[1:])

    add_bold_paragraph(doc, title)
    doc.add_paragraph(body)


# -------------------- 1.1 BIG FIVE --------------------

def add_bigfive_section(doc, user_data):

    doc.add_heading(static_texts.PHASE1_1_SECTION_TITLE, level=1)
    doc.add_paragraph(static_texts.PHASE1_1_INTRO_TEXT)

    bigfive_data = user_data.get("bigfive", {})

    for trait, data in bigfive_data.items():

        score = data.get("score")
        total = data.get("total")
        level = data.get("level", "neutral").lower()

        title = getattr(static_texts, f"PHASE1_1_{trait.upper()}_TITLE")

        #  user score/total score
        full_title = f"{title.strip()} | {score}/{total}"

        # bold text
        add_bold_paragraph(doc, full_title)

        trait_text = personality_texts.BIGFIVE_TEXTS[trait].get(level, "")

        for line in trait_text.strip().split("\n"):
            add_paragraph_conditional_indent(doc, line)


# -------------------- 2.0 LOOPBAANANKERS --------------------

def add_loopbaanankers_section(doc, user_data):

    doc.add_heading(static_texts.PHASE2_0_SECTION_TITLE, level=1)
    doc.add_paragraph(static_texts.PHASE2_0_INTRO_TEXT)

    anchors = user_data.get("loopbaanankers", [])

    for i, anchor in enumerate(anchors[:2], 1):

        title = getattr(static_texts, f"PHASE2_0_HOOGSTE_SCORE_{i}_TITLE")
        text = getattr(loopbaanankers_texts, f"LOOPBAANANKERS_{anchor}")

        doc.add_heading(title.strip(), level=2)

        lines = text.strip().split("\n")

        title = lines[0]
        body = "\n".join(lines[1:])

        add_bold_paragraph(doc, title)
        doc.add_paragraph(body)


# -------------------- 2.1 CARRIERE CLUSTERS --------------------

def add_carriereclusters_section(doc, user_data):

    doc.add_heading(static_texts.PHASE2_1_SECTION_TITLE, level=1)
    doc.add_paragraph(static_texts.PHASE2_1_INTRO_TEXT)

    clusters = user_data.get("carriereclusters", [])

    for i, cluster in enumerate(clusters[:2], 1):

        title = getattr(static_texts, f"PHASE2_1_HOOGSTE_SCORE_{i}_TITLE")
        text = getattr(carriereclusters_texts, f"CARRIER_CLUSTERS_{cluster}")

        # Hoogste score
        doc.add_heading(title.strip(), level=2)

        lines = text.strip().split("\n")
        first_line = lines[0]             # Bold
        body_lines = lines[1:]            # Normal, with bullet indent

        add_bold_paragraph(doc, first_line)
        for line in body_lines:
            add_paragraph_conditional_indent(doc, line)


# -------------------- 2.2 CULTUUR ANALYSE --------------------

def add_cultuur_section(doc, user_data):

    doc.add_heading(static_texts.PHASE2_2_SECTION_TITLE, level=1)
    doc.add_paragraph(static_texts.PHASE2_2_INTRO_TEXT)

    cultures = user_data.get("cultures", [])

    for i, culture in enumerate(cultures[:2], 1):

        title = getattr(static_texts, f"PHASE2_2_HOOGSTE_SCORE_{i}_TITLE")
        text = getattr(cultuuranalyse_texts, f"CULTUUR_ANALYSE_{culture}")

        # Hoogste score
        doc.add_heading(title.strip(), level=2)

        lines = text.strip().split("\n")
        first_line = lines[0]             # Bold
        body_lines = lines[1:]            # Normal, with bullet indent

        add_bold_paragraph(doc, first_line)
        for line in body_lines:
            add_paragraph_conditional_indent(doc, line)


# -------------------- 2.3 JCM --------------------

def add_jcm_section(doc, user_data):

    doc.add_heading(static_texts.PHASE2_3_SECTION_TITLE.strip(), level=1)
    doc.add_paragraph(static_texts.PHASE2_3_INTRO_TEXT.strip())

    jcm_data = user_data.get("jcm", {})

    # Definitions
    for attr, (title_var, text_var) in JCM_MAPPING.items():

        title = getattr(static_texts, f"PHASE2_3_{title_var}")
        explanation = getattr(static_texts, f"PHASE2_3_{text_var}")

        add_bold_paragraph(doc, title.strip())
        doc.add_paragraph(explanation.strip())

    # User answers
    for attr, answer in jcm_data.items():

        add_bold_paragraph(doc, attr)
        doc.add_paragraph((answer or "").strip())
